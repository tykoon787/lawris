# from django.test import TestCase
from .models import Template
import pdfplumber
import docx

# file = "/home/tykoon787/projects/lawris/lawris-app/lawris_django/dms/templates/Form_78_refined.pdf"

# def replace_placeholder(doc, placeholder, replacement):
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             print(f"[DEBUG]: Run: {run.text}")
#             if placeholder in run.text:
#                 run.text = run.text.replace(placeholder, replacement)

# def read_file(file_path):
#     template_content = ""
#     with pdfplumber.open(file_path) as pdf_template:
#         for page in pdf_template.pages:
#             page_text = page.extract_text()
#             template_content += page_text + "\n"
#             return template_content


# def run():
#     # doc = docx.Document(file)

#     replacements = {
#         "{TITLE}": 'Petition for Grant of Probate',
#         "{HEADING}": 'In the High Court of {COURT_LOCATION}',
#         "{PETITIONER_NAME}": 'Kwame Nkuruma',
#         "{PETITIONER_ADDRESS_DESCRIPTION}": 'Residing at',
#         "{PETITIONER_ADDRESS}": '123 Main Street, Accra, Ghana',
#         "{DECEASED_NAME}": 'Ngozi Okafor',
#         "{DATE_OF_DECEASED_DEATH}": '15th September',
#         "{YEAR_OF_DECEASED_DEATH}": '2023',
#         "{PLACE_OF_DECEASED_DEATH}": 'Lagos, Nigeria',
#         "{DOMICILED_AREA}": 'Abuja, Nigeria',
#         "{EXECUTOR}": 'Chukwudi Eze',
#         "{PETITIONER_SIGNATURE}": 'K.N',
#         "{WITNESS}": 'Amina Ibrahim',
#         "{WITNESS_SIGNATURE}": 'A.I',
#         "{WITNESS_ADDRESS_DESCRIPTION}": 'Residing at',
#         "{WITNESS_ADDRESS}": '456 Park Avenue, Nairobi, Kenya',
#         "{ADDRESS_OF_SERVICE}": '789 Court Street, Cape Town, South Africa',
#         "{COURT_LOCATION}": 'Lagos, Nigeria' 
#     }

#     generated_content = read_file(file)
#     filled_template = generated_content.format(**replacements)
#     print(filled_template)

    # # Perform replacements
    # for placeholder, replacement in replacements.items():
    #     replace_placeholder(doc, placeholder, replacement)
    
    # # Save the modified document
    # doc.save("Modifed_v6.docx")  # Replace with the desired output file path



# def main():
#     # Load the Word document
#     doc = docx.Document(file)  # Replace with the path to your Word document
    
#     # Define your replacements
#     replacements = {
#         "{TITLE}": 'Petition for Grant of Probate',
#         "{HEADING}": 'In the High Court of {COURT_LOCATION}',
#         "{PETITIONER_NAME}": 'Kwame Nkuruma',
#         "{PETITIONER_ADDRESS_DESCRIPTION}": 'Residing at',
#         "{PETITIONER_ADDRESS}": '123 Main Street, Accra, Ghana',
#         "{DECEASED_NAME}": 'Ngozi Okafor',
#         "{DATE_OF_DECEASED_DEATH}": '15th September',
#         "{YEAR_OF_DECEASED_DEATH}": '2023',
#         "{PLACE_OF_DECEASED_DEATH}": 'Lagos, Nigeria',
#         "{DOMICILED_AREA}": 'Abuja, Nigeria',
#         "{EXECUTOR}": 'Chukwudi Eze',
#         "{PETITIONER_SIGNATURE}": "K.N",
#         "{WITNESS}": 'Amina Ibrahim',
#         "{WITNESS_SIGNATURE}": "A.I",
#         "{WITNESS_ADDRESS_DESCRIPTION}": 'Residing at',
#         "{WITNESS_ADDRESS}": '456 Park Avenue, Nairobi, Kenya',
#         "{ADDRESS_OF_SERVICE}": '789 Court Street, Cape Town, South Africa',
#         "{COURT_LOCATION}": 'Kisumu, Kenya'
# }

#     # Perform replacements
#     for placeholder, replacement in replacements.items():
#         replace_placeholder(doc, placeholder, replacement)
    
#     # Save the modified document
#     doc.save("Modifed_v6.docx")  # Replace with the desired output file path

# if __name__ == "__main__":
#     main()





# Change this line to reflect the file on your file system

file = "/home/sakwa/Okestra/lawris/lawris-app/lawris_django/dms/templates/Form_78_refined_v2.docx"

form_78_data = {
    'type': 'Petition',
    'title': 'Grant of probate',
    'category_of_law': 'civil',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}

form_79_data = {
    'type': 'Affidavit of Service',
    'title': 'Grant of probate',
    'category_of_law': 'Criminal',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}

form_80_data = {
    'type': 'Amended Plaint',
    'title': 'Grant of probate',
    'category_of_law': 'Civil',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}

form_80_data = {
    'type': 'Grant of Probate Intestate',
    'title': 'Grant of probate',
    'category_of_law': 'Civil',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}

form_81_data = {
    'type': 'Affidavit of Service',
    'title': 'Grant of probate',
    'category_of_law': 'Civil',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}

form_82_data = {
    'type': 'Plaint Fast Track',
    'title': 'Grant of probate',
    'category_of_law': 'Civil',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}

form_83_data = {
    'type': 'Amended Plaint',
    'title': 'Grant of probate',
    'category_of_law': 'Criminal',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}

form_84_data = {
    'type': 'Amended Plaint',
    'title': 'Grant of probate',
    'category_of_law': 'Criminal',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}

form_85_data = {
    'type': 'Grant of Probate Intestate',
    'title': 'Grant of probate',
    'category_of_law': 'Civil',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}

form_86_data = {
    'type': 'Amended Plaint',
    'title': 'Grant of probate',
    'category_of_law': 'Crimial',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}

form_87_data = {
    'type': 'Grant of Probate Intestate',
    'title': 'Grant of probate',
    'category_of_law': 'Civil',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}

form_88_data = {
    'type': 'Plaint Fast Track',
    'title': 'Grant of probate',
    'category_of_law': 'Criminal',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}

form_89_data = {
    'type': 'Amended Plaint',
    'title': 'Grant of probate',
    'category_of_law': 'Criminal',
    'division_of_law': 'family',
    'sub_division': 'succession',
    'template_file_docx': file,
    'pdf_preview_file' : "/mnt/preview",
    'form_fields' : {
        'field1': 'Field 1',
        'field2': 'Field 2',
    }
}
replacements = {
    "TITLE": 'Petition for Grant of Probate',
    "HEADING": 'In the High Court of Nairobi, Kenya',
    "PETITIONER_NAME": 'Kwame Nkuruma',
    "PETITIONER_ADDRESS_DESCRIPTION": 'Residing at',
    "PETITIONER_ADDRESS": '123 Main Street, Nairobi, Kenya',
    "DECEASED_NAME": 'Ngozi Okafor',
    "DATE_OF_DECEASED_DEATH": '15th September',
    "YEAR_OF_DECEASED_DEATH": '2023',
    "PLACE_OF_DECEASED_DEATH": 'Lagos, Nigeria',
    "DOMICILED_AREA": 'Abuja, Nigeria',
    "EXECUTOR": 'Chukwudi Eze',
    "PETITIONER_SIGNATURE": 'K.N',
    "WITNESS_NAME": 'Amina Ibrahim',
    "WITNESS_SIGNATURE": 'A.I',
    "WITNESS_ADDRESS_DESCRIPTION": 'Residing at',
    "WITNESS_ADDRESS": '456 Park Avenue, Nairobi, Kenya',
    "ADDRESS_OF_SERVICE": '789 Court Street, Nairobi, Kenya',
    "COURT_LOCATION": 'Nairobi, Kenya'
}

# replacements = {
#     "{TITLE}": 'Petition for Grant of Probate',
#     "{HEADING}": 'In the High Court of {COURT_LOCATION}',
#     "{PETITIONER_NAME}": 'Kwame Nkuruma',
#     "{PETITIONER_ADDRESS_DESCRIPTION}": 'Residing at',
#     "{PETITIONER_ADDRESS}": '123 Main Street, Accra, Ghana',
#     "{DECEASED_NAME}": 'Ngozi Okafor',
#     "{DATE_OF_DECEASED_DEATH}": '15th September',
#     "{YEAR_OF_DECEASED_DEATH}": '2023',
#     "{PLACE_OF_DECEASED_DEATH}": 'Lagos, Nigeria',
#     "{DOMICILED_AREA}": 'Abuja, Nigeria',
#     "{EXECUTOR}": 'Chukwudi Eze',
#     "{PETITIONER_SIGNATURE}": 'K.N',
#     "{WITNESS}": 'Amina Ibrahim',
#     "{WITNESS_SIGNATURE}": 'A.I',
#     "{WITNESS_ADDRESS_DESCRIPTION}": 'Residing at',
#     "{WITNESS_ADDRESS}": '456 Park Avenue, Nairobi, Kenya',
#     "{ADDRESS_OF_SERVICE}": '789 Court Street, Cape Town, South Africa',
#     "{COURT_LOCATION}": 'Lagos, Nigeria' 
# }

# def replace_placeholder(doc, placeholder, replacement):
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             print(f"[DEBUG]: Run: {run.text}")
#             if placeholder in run.text:
#                 run.text = run.text.replace(placeholder, replacement)

# def run():
#     # form_78 = Template.create_template(**form_78_data)
#     doc = docx.Document(file)

#     # Perform replacements
#     for placeholder, replacement in replacements.items():
#         replace_placeholder(doc, placeholder, replacement)

#     # Save the modified document
#     doc.save("Modifed_v7.docx")  # Replace with the desired output file path


 

#     # doc = docx.Document(file)
#     # filled_document = form_78.fill_template(data=None, document=doc, replacements=replacements)
#     # print(filled_document)

# Added this comment just to test


def main():
    form_78 = Template.create_template(**form_78_data)
    # doc = docx.Document(file)

    # form_78_document = form_78.fill_template(data=None, document=doc, replacements=replacements)
    # print(form_78_document)
    form_79 = Template.create_template(**form_79_data)
    form_80 = Template.create_template(**form_80_data)
    form_81 = Template.create_template(**form_81_data)
    form_82 = Template.create_template(**form_82_data)
    form_83 = Template.create_template(**form_83_data)
    form_84 = Template.create_template(**form_84_data)
    form_85 = Template.create_template(**form_85_data)
    form_86 = Template.create_template(**form_86_data)
    form_87 = Template.create_template(**form_87_data)
    form_88 = Template.create_template(**form_88_data)
    form_89 = Template.create_template(**form_89_data)
def run():
    main()
