from django.db import models
from django.db.models import JSONField
from django.utils import timezone
import uuid
import pdfplumber
import docx
import os
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
import logging
from .utils.microsoft_graph import upload_file_to_onedrive
from bson.binary import Binary, UuidRepresentation
from django.conf import settings


# Logging configuration
log_dir = "/home/kibe/Documents/lawris/logs/"
os.makedirs(log_dir, exist_ok=True)

logger = logging.getLogger()
logger.setLevel(logging.DEBUG)

log_file = os.path.join(log_dir, "tests.py.log")
file_handler = logging.FileHandler(log_file, mode="w")
file_handler.setLevel(logging.DEBUG)
file_formatter = logging.Formatter(
    '[%(asctime)s] [%(levelname)-5s]  ::    %(message)s')
file_handler.setFormatter(file_formatter)
logger.addHandler(file_handler)


# MongoDB connection
# Access the MongoDB client
mongodb_client = settings.MONGODB_CLIENT

# Access the MongoDB database and collection
mongodb_database = settings.MONGODB_DATABASE
mongodb_collection = settings.MONGODB_COLLECTION


class BaseModel(models.Model):
    """
    Base Model for all models
    """
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    created_at = models.DateTimeField(default=timezone.now, editable=False)
    updated_at = models.DateTimeField(default=timezone.now, editable=False)

    class Meta:
        abstract = True

    def save(self, *args, **kwargs):
        """
        Override the save method to update the 'upadated_at' timestamp when saving an instance
        """

        self.updated_at = timezone.now()
        super(BaseModel, self).save(*args, **kwargs)


class Template(BaseModel):
    """
    Template Model
    """
    TEMPLATE_TYPES = [
        ('affidavit', 'Affidavit'),
        ('petition', 'Petition'),
        ('plaint', 'Plaint'),
        ('notice of motion', 'Notice of Motion'),
        ('defense', 'Defense'),
        ('counter claim', 'Counter Claim'),
        ('judgment', 'Judgment')
    ]

    CATEGORY_CHOICES = [
        ('civil', 'Civil'),
        ('criminal', 'Criminal'),
        ('commercial', 'Commercial'),
        ('land', 'Land'),
        ('arbitration', 'Arbitration'),
        ('family', 'Family'),
        ('business', 'Business'),
        ('property', 'Property'),
    ]

    DIVISION_CHOICES = [
        ('family', 'Family'),
        ('land', 'Land'),
        ('employment', 'Employment'),
        ('tax', 'Tax'),
        ('intellectual property', 'Intellectual Property'),
    ]

    SUB_DIVISION_CHOICES = [
        ('succession', 'Succession'),
        ('divorce', 'Divorce'),
        ('real estate', 'Real Estate'),
        ('labor', 'Labor'),
        ('patent', 'Patent'),
    ]

    type = models.CharField(max_length=20, choices=TEMPLATE_TYPES)
    title = models.CharField(max_length=200)
    category_of_law = models.CharField(max_length=20, choices=CATEGORY_CHOICES)
    division_of_law = models.CharField(max_length=50, choices=DIVISION_CHOICES)
    sub_division = models.CharField(
        max_length=20, choices=SUB_DIVISION_CHOICES)
    template_file_docx = models.CharField(max_length=255)
    pdf_preview_file = models.CharField(max_length=255, default="preview_file")
    form_fields = JSONField(blank=True, null=True)

    @classmethod
    def create_template(cls, type, title, category_of_law, division_of_law, sub_division, template_file_docx, pdf_preview_file, form_fields):
        """
        Creates a new template and stores it in the database

        Args:
            type (str): The type of the template
            title (str): The title of the template
            category_of_law (str): The category of law
            division_of_law (str): The division of law
            sub_division (str): The branch of division of law
            template_file_docx (str): The path to the template file in 'docx' format
            pdf_preview_file (str) : The path to the pdf preview file
            thumbnail (str) : The path to the thumbnail of the file
            form_fields (dict): A JSON dictionary representing form fields

        Returns:
            Template: The newly created template instance
        """
        template = cls(
            type=type,
            title=title,
            category_of_law=category_of_law,
            division_of_law=division_of_law,
            sub_division=sub_division,
            template_file_docx=template_file_docx,
            pdf_preview_file=pdf_preview_file,
            #thumbnail=thumbnail,
            form_fields=form_fields
        )
        template.save()
        return template

    def read_template_content(self, template_file: str):
        """
        Reads the content from a template file

        Args: 
            template_file (str): Path to the file
        """
        template_content = ""
        with pdfplumber.open(template_file) as pdf_template:
            for page in pdf_template.pages:
                page_text = page.extract_text()
                template_content += page_text + "\n"
        print(template_content)
        return template_content

    def replace_placeholder(self, doc: docx.Document, placeholder: str, replacement: str):
        """
        Replaces a placeholder from a ``.docx`` document with the
        replacement string
        """
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(
                        placeholder, replacement.upper())

    def save_file(self, document: docx.Document, file_name: str = "modified.docx", save_path: str = "/home/kibe/Documents/lawris/docs/"):
        """
        Saves the generated file to the specified path 

        Args:
            doc: The document to be saved
            file_name: The file name to use
            save_path: The path where the file is to be saved

        Return:
            Nothing
        """
        file_name = os.path.join(save_path, file_name)
        document.save(file_name)
        
        # Upload the file to OneDrive
        # upload_file_to_onedrive(save_path, file_name)
        with open(file_name, 'rb') as file:
            generated_content = file.read()

        return generated_content

    def fill_template(self, data=None, document: docx.Document = None, replacements: dict = None):
        """
        Fills the template with data
        If ``docx`` is passed, the placeholders are replaced within te word document

        Args: 
            data (dict): Data to be filled
            template_file: Template file to be read and used
            document (docx.Document): Document to be read and replaced
            replacements (dict) : Dictionary containing the placeholders and replacement

        Returns:
            Document (Document): Document object with the generated content
        """
        
        
        if (document and replacements):
            for placeholder, replacement in replacements.items():
                self.replace_placeholder(document, placeholder, replacement)

            new_document = Document(title=self.title)
            
            document_id = new_document.id
            # Convert UUID to binary with specified UuidRepresentation
            
            document_id_binary = Binary.from_uuid(document_id, UuidRepresentation.STANDARD)


            template_data = {
                "type": self.type,
                "title": self.title,
                "category_of_law": self.category_of_law,
                "division_of_law": self.division_of_law,
                "sub_division": self.sub_division,
                "document_location": "http//:onedrive.com",
                "thumbnail_url": "thumbnail location",
                "document_id": document_id_binary
                
                
            }
            mongodb_collection.insert_one(template_data)

            generated_content = self.save_file(
                new_document, file_name=f'{new_document.title}.docx')
            new_document.content = generated_content       
            return new_document
        
        else:
            template_content = self.read_template_content(
                self.template_file_docx)
            filled_template = template_content.format(**replacements)
            new_document = Document(self.title)
            document_id = new_document.id
            # Convert UUID to binary with specified UuidRepresentation
            document_id_binary = Binary.from_uuid(document_id, UuidRepresentation.STANDARD)
            template_data = {
                "type": self.type,
                "title": self.title,
                "category_of_law": self.category_of_law,
                "division_of_law": self.division_of_law,
                "sub_division": self.sub_division,
                "document_location": "http//:onedrive.com",
                "thumbnail_url": "thumbnail location pdf file",
                "document_id": document_id_binary
            }
            mongodb_collection.insert_one(template_data)
        
            return filled_template

    def __str__(self):
        return self.title


class Document(BaseModel):
    """
    Document Model
    """
    title = models.CharField(max_length=200)
    content = models.TextField(blank=True, null=True)

    def __init__(self, title, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.title = title

    def generate_document(self, format: str = "docx"):
        if format.lower() == "docx":
            doc = docx.Document()
            doc.add_paragraph(str(self.content))
            print("Generated document:")
            print(doc.paragraphs)  # Log the paragraphs
            return doc
        else:
            # Handle other formats (pdf)
            pass

    def download_document(self):
        generated_doc = self.generate_document()

        if generated_doc:
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename={self.title}.docx'
            generated_doc.save(response)

            return response